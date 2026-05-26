from __future__ import annotations

import io

import pytest

from app.services.bdd_parsers import ParseError, parse_source


# ---------------------------------------------------------------------------
# 1. Text passthrough
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_text_passthrough():
    result = await parse_source("text", content="Hello requirement", url=None)
    assert result == "Hello requirement"


# ---------------------------------------------------------------------------
# 2. Text with empty content raises ParseError
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_text_empty_raises():
    with pytest.raises(ParseError, match="content"):
        await parse_source("text", content=None, url=None)


# ---------------------------------------------------------------------------
# 3. Unknown source_type raises ParseError
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_unknown_source_raises():
    with pytest.raises(ParseError, match="unknown"):
        await parse_source("fax", content=None, url=None)


# ---------------------------------------------------------------------------
# 4. URL source with missing url raises ParseError
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_url_missing_raises():
    with pytest.raises(ParseError, match="url"):
        await parse_source("url", content=None, url=None)


# ---------------------------------------------------------------------------
# 5. Confluence without credentials raises ParseError
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_confluence_no_credentials_raises():
    with pytest.raises(ParseError, match="credentials"):
        await parse_source(
            "confluence",
            content=None,
            url="https://example.atlassian.net/wiki/pages/123",
        )


# ---------------------------------------------------------------------------
# 6. PDF round-trip using pypdf
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_pdf_bytes():
    import pypdf
    from pypdf import PdfWriter

    buf = io.BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=200, height=200)

    # Add text to the page so extract_text returns something
    page.merge_page(page)  # no-op merge but keeps object valid

    # Use a simpler approach: write a real PDF and check we can call the parser
    # without errors (blank page may yield empty string — we verify no crash)
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    # Should not raise
    result = await parse_source("pdf", content=None, url=None, file_bytes=pdf_bytes)
    assert isinstance(result, str)


# ---------------------------------------------------------------------------
# 7. DOCX round-trip using python-docx
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_parse_docx_bytes():
    import docx

    doc = docx.Document()
    doc.add_paragraph("First requirement paragraph.")
    doc.add_paragraph("Second requirement paragraph.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = await parse_source("docx", content=None, url=None, file_bytes=docx_bytes)
    assert "First requirement paragraph." in result
    assert "Second requirement paragraph." in result
